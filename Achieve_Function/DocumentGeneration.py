from docx import Document
import docx

# 输入数据
expression = {'anger': 0, 'disgust': 0, 'fear': 0, 'happy': 10, 'sad': 0, 'surprised': 0, 'normal': 0,'Noface':0}
posture = {'cross_arms': 0, 'shoulder_drop': 0}


def getMaxExpression(expression):
    maxnum = 0;
    maxkey = ''
    for keys, value in expression.items():
        if value > maxnum:
            maxnum = value
            maxkey = keys
    return maxkey


def generateDocument(expression, posture, total):
    document = Document()



    # 添加异常姿态表格
    document.add_heading('田昊 学习状态评估', level=2)
    document.add_heading('非放松自然状态动作统计', level=2)
    posture_table = document.add_table(rows=2, cols=2)
    row = posture_table.rows[0]
    row.cells[0].text = '环抱双臂动作'
    row.cells[1].text = str(posture['cross_arms'])

    row = posture_table.rows[1]
    row.cells[0].text = '肩部高低不一致'
    row.cells[1].text = str(posture['shoulder_drop'])

    # 添加表情统计表格
    document.add_heading('面部表情状态数据统计', level=2)
    table = document.add_table(rows=7, cols=2)
    row = table.rows[0]
    row.cells[0].text = 'anger'
    row.cells[1].text = str(expression['anger'])

    row = table.rows[1]
    row.cells[0].text = 'disgust'
    row.cells[1].text = str(expression['disgust'])

    row = table.rows[2]
    row.cells[0].text = 'fear';
    row.cells[1].text = str(expression['fear'])

    row = table.rows[3]
    row.cells[0].text = 'happy'
    row.cells[1].text = str(expression['happy'])

    row = table.rows[4]
    row.cells[0].text = 'sad'
    row.cells[1].text = str(expression['sad'])

    row = table.rows[5]
    row.cells[0].text = 'surprised'
    row.cells[1].text = str(expression['surprised'])

    row = table.rows[6]
    row.cells[0].text = 'normal'
    row.cells[1].text = str(expression['normal'])



    document.add_heading('学习状态评估结果', level=2)

    # 总结出现次数最多的表情  对各个表情种类的出现进行分析

    maxExpression = getMaxExpression(expression);

    # 总结 normal 和happy 出现占比
    study_state_text = ''
    normal_happy = (expression['normal'] + expression['happy']) / total*100

    Noface = (expression['Noface']) / total*100



    # 总结异常姿势出现 进行相应的提醒
    # posture_text =''
    cross_arms_text = ''
    shoulder_drop_text = ''
    Noface_text=''
    if posture['cross_arms'] > 0:
        cross_arms_text = '出现双臂环绕动作，易导致出现紧张不适心理状态，影响听课状态，建议放松双臂，保持放松坐姿'
        normal_happy-10
    if posture['shoulder_drop'] > 0:
        shoulder_drop_text = '两侧肩部高度差过大，建议保持肩膀放松且平衡'
        normal_happy-10
    if normal_happy >= 80:
        study_state_text = '非常良好'
    if normal_happy > 70 and normal_happy <=80:
        study_state_text = '良好'
    if normal_happy > 50 and normal_happy <70:
        study_state_text = '一般'
    if normal_happy <= 50:
        study_state_text = '较差'
    if Noface >=30:
        study_state_text='较差'
        Noface_text='出现没有识别到面部的情况，有可能存在逃课行为'

    paragraph1 = document.add_paragraph("学生学习状态评价："+study_state_text)
    paragraph2 = document.add_paragraph("出现最多的表情为：" + maxExpression+"占整个学习阶段的："+str(normal_happy)+"%")
    paragraph3 = document.add_paragraph(cross_arms_text)
    paragraph3 = document.add_paragraph(shoulder_drop_text)
    paragraph4 = document.add_paragraph(Noface_text)

    document.save('a.docx')
    print(total)

#generateDocument(expression, posture, 100)
